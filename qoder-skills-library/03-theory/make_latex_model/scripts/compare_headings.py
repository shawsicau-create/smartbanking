#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标题文字对比工具
对比“基准模板（推荐：PDF）”与 LaTeX 文件的标题文字差异

使用方法:
    # 对比两个文件
    python scripts/compare_headings.py baseline.pdf main.tex

    # 输出为 HTML 报告
    python scripts/compare_headings.py baseline.pdf main.tex --report output.html

    # 输出为 Markdown 报告
    python scripts/compare_headings.py baseline.pdf main.tex --report output.md
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple
from datetime import datetime
import warnings

# 允许在任何 cwd 下运行时都能导入同目录脚本
SCRIPT_DIR = Path(__file__).parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))


def extract_from_latex(tex_file: Path, check_format: bool = False) -> Dict[str, any]:
    """
    从 LaTeX 文件中提取标题文字

    Args:
        tex_file: LaTeX 文件路径
        check_format: 是否检查格式（加粗）

    Returns:
        如果 check_format=False: Dict[str, str] - 标题文本
        如果 check_format=True: Dict[str, Dict] - 包含文本和格式信息
    """
    headings = {}

    with open(tex_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 预处理：去掉注释（避免把被注释掉的 \subsection{...} 误识别为标题）
    content_no_comments_lines = []
    for line in content.splitlines():
        cleaned = re.sub(r"(?<!\\)%.*$", "", line)
        content_no_comments_lines.append(cleaned)
    content_no_comments = "\n".join(content_no_comments_lines)

    def _extract_braced_arg(src: str, brace_start_idx: int) -> Tuple[str, int]:
        """从 src[brace_start_idx] == '{' 开始提取配对花括号内容（支持嵌套）。"""
        if brace_start_idx < 0 or brace_start_idx >= len(src) or src[brace_start_idx] != "{":
            return "", brace_start_idx

        depth = 1
        i = brace_start_idx + 1
        arg_start = i

        while i < len(src) and depth > 0:
            ch = src[i]
            # 跳过转义字符（避免把 \{ \} 误判为结构花括号）
            if ch == "\\" and i + 1 < len(src):
                i += 2
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
            i += 1

        if depth != 0:
            # 括号不平衡，回退为空
            return "", brace_start_idx

        return src[arg_start : i - 1], i

    def _iter_command_args(src: str, command: str):
        # 允许命令与 { 之间有空白
        for m in re.finditer(rf"\\{re.escape(command)}\s*\{{", src):
            brace_idx = m.end() - 1  # 指向 '{'
            arg, end_idx = _extract_braced_arg(src, brace_idx)
            if arg:
                yield (m.start(), arg)

    # 以“文档顺序”为准同时提取 \section{} 与（NSFC 模板常用的）\NSFCSubsection{} / \subsection{}
    tokens = []
    for pos, arg in _iter_command_args(content_no_comments, "section"):
        tokens.append((pos, "section", arg))
    for pos, arg in _iter_command_args(content_no_comments, "NSFCSubsection"):
        tokens.append((pos, "nsfc_subsection", arg))
    for pos, arg in _iter_command_args(content_no_comments, "subsection"):
        tokens.append((pos, "subsection", arg))

    tokens.sort(key=lambda x: x[0])

    section_num = 0
    subsection_num = 0

    for _, kind, raw in tokens:
        if kind == "section":
            section_num += 1
            subsection_num = 0
            key = f"section_{section_num}"
        else:
            if section_num <= 0:
                # 忽略“没有 section 上下文”的二级标题（通常不应出现）
                continue
            subsection_num += 1
            key = f"subsection_{section_num}_{subsection_num}"

        if check_format:
            headings[key] = {
                "text": clean_latex_text(raw),
                "fragments": extract_formatted_text_from_latex(raw),
            }
        else:
            headings[key] = clean_latex_text(raw)

    return headings


def clean_latex_text(text: str) -> str:
    """清理 LaTeX 文本中的格式标记"""
    try:
        from core.latex_format_parser import LatexFormatParser

        cleaned = LatexFormatParser.clean_latex_text(text)
        cleaned = cleaned.replace("~", " ")
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        return cleaned
    except Exception:
        # fallback: 旧版正则（不支持嵌套/声明式格式）
        text = re.sub(r'\\[a-zA-Z]+', '', text)
        text = re.sub(r'\{|\}', '', text)
        # 渲染层面的空白归一：~ 在 TeX 中等价于不换行空格
        text = text.replace('~', ' ')
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text


def clean_latex_commands(text: str) -> str:
    """清理 LaTeX 命令，但保留 \textbf 和 \bfseries"""
    # 删除除 \textbf、\bfseries 外的所有命令
    text = re.sub(r'\\(?!textbf|bfseries)[a-zA-Z]+', '', text)
    text = re.sub(r'\{|\}', '', text)
    # 渲染层面的空白归一：~ 在 TeX 中等价于不换行空格
    text = text.replace('~', ' ')
    # 格式对比需要保留片段边界处的空格（例如 "1. "），因此不做 strip()。
    text = re.sub(r'\s+', ' ', text)
    return text


def extract_formatted_text_from_word(paragraph) -> List[Dict[str, any]]:
    """
    从 Word 段落中提取带格式信息的文本片段

    Args:
        paragraph: python-docx 的段落对象

    Returns:
        [
            {"text": "立项依据", "bold": True},
            {"text": "与研究内容", "bold": False}
        ]
    """
    fragments = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        fragments.append({
            "text": text,
            "bold": run.bold if run.bold is not None else False
        })
    return fragments


def extract_formatted_text_from_latex(latex_text: str) -> List[Dict[str, any]]:
    """
    从 LaTeX 文本中提取带格式信息的片段

    支持的格式：
    - \textbf{文本}     （推荐）
    - {\bfseries 文本}  （传统方式）

    Args:
        latex_text: LaTeX 标题文本

    Returns:
        [
            {"text": "立项依据", "bold": True},
            {"text": "与研究内容", "bold": False}
        ]
    """
    try:
        from core.latex_format_parser import LatexFormatParser

        parsed = LatexFormatParser.extract_formatted_text(latex_text)
        # compare_headings 的格式对比目前只关心 bold；但保留其它字段不影响结果
        out: List[Dict[str, any]] = []
        for frag in parsed:
            t = str(frag.get("text", "") or "")
            t = t.replace("\u00a0", " ").replace("~", " ")
            if not t:
                continue
            out.append({"text": t, "bold": bool(frag.get("bold", False))})

        # 合并相邻同样 bold 的片段，减少字符级对齐噪声
        merged: List[Dict[str, any]] = []
        for frag in out:
            if not merged:
                merged.append(frag)
                continue
            if bool(merged[-1].get("bold")) == bool(frag.get("bold")):
                merged[-1]["text"] += frag["text"]
            else:
                merged.append(frag)
        return merged
    except Exception:
        # fallback: 旧版只识别 \textbf{...}
        fragments: List[Dict[str, any]] = []
        textbf_pattern = r'\\textbf\{([^}]+)\}'
        bold_segments = []
        for match in re.finditer(textbf_pattern, latex_text):
            bold_segments.append(
                {"start": match.start(), "end": match.end(), "text": match.group(1), "bold": True}
            )
        bold_segments.sort(key=lambda x: x["start"])
        last_end = 0
        for seg in bold_segments:
            if seg["start"] > last_end:
                normal_text = latex_text[last_end:seg["start"]]
                normal_text = clean_latex_commands(normal_text)
                if normal_text:
                    fragments.append({"text": normal_text, "bold": False})
            fragments.append({"text": seg["text"].replace("~", " "), "bold": True})
            last_end = seg["end"]
        if last_end < len(latex_text):
            normal_text = latex_text[last_end:]
            normal_text = clean_latex_commands(normal_text)
            if normal_text:
                fragments.append({"text": normal_text, "bold": False})
        return fragments


def compare_formatted_text(word_fragments: List[Dict],
                          latex_fragments: List[Dict]) -> Dict[str, any]:
    """
    对比 Word 和 LaTeX 的格式化文本

    Args:
        word_fragments: Word 格式片段列表
        latex_fragments: LaTeX 格式片段列表

    Returns:
        {
            "match": true/false,
            "word_text": "立项依据与研究内容",
            "latex_text": "立项依据与研究内容",
            "differences": [
                {
                    "type": "bold_mismatch",
                    "word_fragment": {"text": "立项依据", "bold": True},
                    "latex_fragment": {"text": "立项依据", "bold": False},
                    "position": "0-4"
                }
            ]
        }
    """
    # 提取纯文本进行初步对比
    word_text = "".join(f["text"] for f in word_fragments)
    latex_text = "".join(f["text"] for f in latex_fragments)

    if word_text != latex_text:
        return {
            "match": False,
            "reason": "text_mismatch",
            "word_text": word_text,
            "latex_text": latex_text,
            "word_fragments": word_fragments,
            "latex_fragments": latex_fragments,
        }

    # 对齐片段并对比格式
    differences = []
    word_pos = 0
    word_idx = 0
    latex_idx = 0

    # 创建可修改的片段副本
    word_frags = [f.copy() for f in word_fragments]
    latex_frags = [f.copy() for f in latex_fragments]

    while word_idx < len(word_frags) and latex_idx < len(latex_frags):
        word_frag = word_frags[word_idx]
        latex_frag = latex_frags[latex_idx]

        # 计算当前片段的文本长度
        word_len = len(word_frag["text"])
        latex_len = len(latex_frag["text"])

        # 找到最小长度
        min_len = min(word_len, latex_len)

        # 对比前 min_len 个字符的格式
        for i in range(min_len):
            if word_frag["bold"] != latex_frag["bold"]:
                char_pos = word_pos + i
                differences.append({
                    "type": "bold_mismatch",
                    "position": char_pos,
                    "char": word_frag["text"][i],
                    "word_bold": word_frag["bold"],
                    "latex_bold": latex_frag["bold"]
                })

        # 更新位置
        word_pos += min_len
        word_frag["text"] = word_frag["text"][min_len:]
        latex_frag["text"] = latex_frag["text"][min_len:]
        word_len -= min_len
        latex_len -= min_len

        # 如果 Word 片段用完了，移到下一个
        if word_len == 0:
            word_idx += 1
        # 如果 LaTeX 片段用完了，移到下一个
        if latex_len == 0:
            latex_idx += 1

    return {
        "match": len(differences) == 0,
        "word_text": word_text,
        "latex_text": latex_text,
        "differences": differences,
        "word_fragments": word_fragments,
        "latex_fragments": latex_fragments,
    }


def extract_from_word(doc_file: Path, check_format: bool = False) -> Dict[str, any]:
    """
    ⚠️ 兼容保留：建议改用 PDF 作为标题/格式基准。

    从 Word 文档中提取标题文字

    Args:
        doc_file: Word 文档路径
        check_format: 是否检查格式（加粗）

    Returns:
        如果 check_format=False: Dict[str, str] - 标题文本
        如果 check_format=True: Dict[str, Dict] - 包含文本和格式信息
    """
    warnings.warn(
        "Word(.docx) 标题提取仅为向后兼容保留；推荐使用 PDF 基准（Single Source of Truth）。",
        DeprecationWarning,
        stacklevel=2,
    )
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("安装命令: pip install python-docx")
        sys.exit(1)

    if not doc_file.suffix == '.docx':
        print(f"警告: {doc_file} 是 .doc 格式，建议转换为 .docx")
        sys.exit(1)

    doc = Document(doc_file)

    def _add_heading(out: Dict[str, any], key: str, paragraph):
        if check_format:
            out[key] = {
                "text": paragraph.text.strip(),
                "fragments": extract_formatted_text_from_word(paragraph),
            }
        else:
            out[key] = paragraph.text.strip()

    headings: Dict[str, any] = {}

    # 1) 优先走“标准标题样式”路径（适配常规 Word 文档）
    section_count = 0
    subsection_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading 1" in style_name or "标题 1" in style_name:
            section_count += 1
            subsection_count = 0
            _add_heading(headings, f"section_{section_count}", paragraph)
        elif "Heading 2" in style_name or "标题 2" in style_name:
            if section_count <= 0:
                continue
            subsection_count += 1
            _add_heading(headings, f"subsection_{section_count}_{subsection_count}", paragraph)

    if headings:
        return headings

    # 2) 回退：NSFC 等模板常把提纲标题设为 Normal 样式（用文本模式识别）
    section_re = re.compile(r"^（[一二三四五六七八九十]+）")
    subsection_re = re.compile(r"^\s*\d+\s*[\.．、]")

    section_count = 0
    subsection_count = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if section_re.match(text):
            section_count += 1
            subsection_count = 0
            _add_heading(headings, f"section_{section_count}", paragraph)
            continue

        if subsection_re.match(text):
            if section_count <= 0:
                continue
            subsection_count += 1
            _add_heading(headings, f"subsection_{section_count}_{subsection_count}", paragraph)

    return headings


def extract_from_pdf(pdf_file: Path, check_format: bool = False) -> Dict[str, any]:
    """从 PDF 中提取标题文字（可选：加粗片段、换行点）。"""
    try:
        from extract_headings_from_pdf import extract_headings_from_pdf
    except Exception as e:
        raise RuntimeError(f"无法导入 extract_headings_from_pdf.py: {e}")
    return extract_headings_from_pdf(pdf_file, check_format=check_format)


def extract_from_source(source_file: Path, check_format: bool = False) -> Dict[str, any]:
    """
    从基准源提取标题（推荐：PDF；兼容：DOCX）。

    优先级：
    - .pdf  → extract_from_pdf
    - .docx → extract_from_word（deprecated）
    """
    suf = source_file.suffix.lower()
    if suf == ".pdf":
        return extract_from_pdf(source_file, check_format=check_format)
    if suf == ".docx":
        return extract_from_word(source_file, check_format=check_format)
    raise ValueError(f"不支持的基准文件格式: {source_file}")


def compare_headings(word_headings: Dict[str, str], latex_headings: Dict[str, str]) -> Tuple[List, List, List]:
    """
    对比两个标题字典（仅文本对比）

    Returns:
        (完全匹配的列表, 有差异的列表, 仅在一方存在的列表)
    """
    all_keys = set(word_headings.keys()) | set(latex_headings.keys())

    matched = []
    differences = []
    only_in_one = []

    for key in sorted(all_keys):
        word_value = word_headings.get(key, '')
        latex_value = latex_headings.get(key, '')

        if word_value == latex_value:
            if word_value:  # 两者都有且相同
                matched.append((key, word_value))
        else:
            if word_value and latex_value:  # 两者都有但不同
                differences.append((key, word_value, latex_value))
            elif word_value:  # 仅在 Word 中
                only_in_one.append(('word', key, word_value))
            elif latex_value:  # 仅在 LaTeX 中
                only_in_one.append(('latex', key, latex_value))

    return matched, differences, only_in_one


def compare_headings_with_format(word_headings: Dict[str, Dict],
                                 latex_headings: Dict[str, Dict]) -> Tuple[List, List, List, List]:
    """
    对比两个标题字典（包含格式对比）

    Returns:
        (完全匹配的列表, 文本差异列表, 格式差异列表, 仅在一方存在的列表)
    """
    all_keys = set(word_headings.keys()) | set(latex_headings.keys())

    matched = []
    text_diff = []
    format_diff = []
    only_in_one = []

    for key in sorted(all_keys):
        word_data = word_headings.get(key)
        latex_data = latex_headings.get(key)

        if not word_data and not latex_data:
            continue

        if not word_data:
            only_in_one.append(('latex', key, latex_data["text"]))
        elif not latex_data:
            only_in_one.append(('word', key, word_data["text"]))
        else:
            # 两者都存在，对比文本和格式
            word_text = word_data["text"]
            latex_text = latex_data["text"]

            if word_text != latex_text:
                # 文本不一致
                text_diff.append((key, word_text, latex_text))
            else:
                # 文本一致，对比格式
                format_result = compare_formatted_text(
                    word_data["fragments"],
                    latex_data["fragments"]
                )

                if format_result["match"]:
                    matched.append((key, word_text, format_result))
                else:
                    format_diff.append((key, word_text, format_result))

    return matched, text_diff, format_diff, only_in_one


def generate_text_report_with_format(matched: List, text_diff: List, format_diff: List, only_in_one: List) -> str:
    """生成文本格式报告（包含格式对比）"""
    lines = []
    lines.append('=' * 60)
    lines.append('  标题文字对比报告（包含格式）')
    lines.append('=' * 60)
    lines.append('')

    # 统计
    total = len(matched) + len(text_diff) + len(format_diff)
    match_count = len(matched)
    text_diff_count = len(text_diff)
    format_diff_count = len(format_diff)
    only_count = len(only_in_one)

    lines.append(f'总标题数: {total}')
    lines.append(f'✅ 完全匹配（文本+格式）: {match_count}')
    lines.append(f'⚠️  文本差异: {text_diff_count}')
    lines.append(f'🔶 格式差异: {format_diff_count}')
    lines.append(f'❌ 仅在一方: {only_count}')
    lines.append('')

    # 完全匹配的标题
    if matched:
        lines.append('# 完全匹配的标题')
        lines.append('')
        for key, value, _ in matched:
            lines.append(f'✅ {key}: {value}')
        lines.append('')

    # 文本差异
    if text_diff:
        lines.append('# 文本差异')
        lines.append('')
        for key, word_value, latex_value in text_diff:
            lines.append(f'⚠️  {key}:')
            lines.append(f'   Word:  {word_value}')
            lines.append(f'   LaTeX: {latex_value}')
            lines.append('')

    # 格式差异
    if format_diff:
        lines.append('# 格式差异（加粗）')
        lines.append('')
        for key, text, result in format_diff:
            lines.append(f'🔶 {key}: {text}')
            lines.append('   格式差异:')

            # 显示 Word 格式
            word_display = []
            for frag in result.get("word_fragments", []):
                marker = '**' if frag["bold"] else ''
                word_display.append(f'{marker}{frag["text"]}{marker}')
            lines.append(f'   Word:  {"".join(word_display)}')

            # 显示 LaTeX 格式
            latex_display = []
            for frag in result.get("latex_fragments", []):
                marker = '**' if frag["bold"] else ''
                latex_display.append(f'{marker}{frag["text"]}{marker}')
            lines.append(f'   LaTeX: {"".join(latex_display)}')

            # 显示差异详情
            if result.get("differences"):
                lines.append('   差异位置:')
                for diff in result["differences"]:
                    char = diff.get("char", "")
                    word_bold = "加粗" if diff.get("word_bold") else "正常"
                    latex_bold = "加粗" if diff.get("latex_bold") else "正常"
                    lines.append(f'     位置 {diff.get("position")}: "{char}" - Word:{word_bold}, LaTeX:{latex_bold}')
            lines.append('')

    # 仅在一方的标题
    if only_in_one:
        lines.append('# 仅在一方的标题')
        lines.append('')
        for source, key, value in only_in_one:
            source_label = 'Word' if source == 'word' else 'LaTeX'
            lines.append(f'❌ 仅在 {source_label}: {key}')
            lines.append(f'   {value}')
            lines.append('')

    return '\n'.join(lines)


def generate_text_report(matched: List, differences: List, only_in_one: List) -> str:
    """生成文本格式报告"""
    lines = []
    lines.append('=' * 60)
    lines.append('  标题文字对比报告')
    lines.append('=' * 60)
    lines.append('')

    # 统计
    total = len(matched) + len(differences)
    match_count = len(matched)
    diff_count = len(differences)
    only_count = len(only_in_one)

    lines.append(f'总标题数: {total}')
    lines.append(f'✅ 完全匹配: {match_count}')
    lines.append(f'⚠️  有差异: {diff_count}')
    lines.append(f'❌ 仅在一方: {only_count}')
    lines.append('')

    # 完全匹配的标题
    if matched:
        lines.append('# 完全匹配的标题')
        lines.append('')
        for key, value in matched:
            lines.append(f'✅ {key}: {value}')
        lines.append('')

    # 有差异的标题
    if differences:
        lines.append('# 有差异的标题')
        lines.append('')
        for key, word_value, latex_value in differences:
            lines.append(f'⚠️  {key}:')
            lines.append(f'   Word:  {word_value}')
            lines.append(f'   LaTeX: {latex_value}')
            lines.append('')

    # 仅在一方的标题
    if only_in_one:
        lines.append('# 仅在一方的标题')
        lines.append('')
        for source, key, value in only_in_one:
            source_label = 'Word' if source == 'word' else 'LaTeX'
            lines.append(f'❌ 仅在 {source_label}: {key}')
            lines.append(f'   {value}')
            lines.append('')

    return '\n'.join(lines)


def generate_html_report(matched: List, differences: List, only_in_one: List,
                         word_file: Path, latex_file: Path) -> str:
    """生成 HTML 格式报告（仅文本对比）"""
    match_count = len(matched)
    diff_count = len(differences)
    only_count = len(only_in_one)

    def _esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>标题文字对比报告</title>
  <style>
    body {{
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
      line-height: 1.6;
      max-width: 1200px;
      margin: 0 auto;
      padding: 20px;
      background: #f5f5f5;
    }}
    .header {{
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white;
      padding: 24px;
      border-radius: 10px;
      margin-bottom: 20px;
      box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }}
    .meta {{
      opacity: 0.9;
      font-size: 14px;
      margin-top: 8px;
    }}
    .stats {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
      gap: 16px;
      margin-bottom: 20px;
    }}
    .stat-card {{
      background: white;
      padding: 16px;
      border-radius: 8px;
      box-shadow: 0 2px 4px rgba(0,0,0,0.1);
      text-align: center;
    }}
    .stat-card h3 {{
      margin: 0 0 8px 0;
      font-size: 13px;
      color: #666;
      font-weight: 600;
    }}
    .stat-card .value {{
      font-size: 28px;
      font-weight: bold;
    }}
    .matched .value {{ color: #10b981; }}
    .differences .value {{ color: #f59e0b; }}
    .only .value {{ color: #ef4444; }}
    .section {{
      background: white;
      padding: 20px;
      border-radius: 8px;
      margin-bottom: 16px;
      box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }}
    .section h2 {{
      margin: 0 0 14px 0;
      padding-bottom: 10px;
      border-bottom: 2px solid #e5e7eb;
      font-size: 18px;
    }}
    .item {{
      padding: 12px;
      margin-bottom: 10px;
      border-left: 4px solid #ddd;
      background: #f9fafb;
      border-radius: 4px;
    }}
    .item.matched {{
      border-left-color: #10b981;
      background: #f0fdf4;
    }}
    .item.difference {{
      border-left-color: #f59e0b;
      background: #fffbeb;
    }}
    .item.only {{
      border-left-color: #ef4444;
      background: #fef2f2;
    }}
    .key {{
      font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace;
      font-size: 13px;
      color: #374151;
      margin-bottom: 6px;
    }}
    .value-block {{
      font-size: 14px;
      margin: 2px 0;
      white-space: pre-wrap;
    }}
    .label {{
      display: inline-block;
      min-width: 52px;
      font-weight: 600;
      color: #111827;
    }}
  </style>
</head>
<body>
  <div class="header">
    <h1 style="margin:0; font-size: 24px;">标题文字对比报告（仅文本）</h1>
    <div class="meta">Word: {_esc(str(word_file))}<br>LaTeX: {_esc(str(latex_file))}</div>
  </div>

  <div class="stats">
    <div class="stat-card matched"><h3>完全匹配</h3><div class="value">{match_count}</div></div>
    <div class="stat-card differences"><h3>有差异</h3><div class="value">{diff_count}</div></div>
    <div class="stat-card only"><h3>仅在一方</h3><div class="value">{only_count}</div></div>
  </div>
"""

    if matched:
        html += '<div class="section"><h2>完全匹配</h2>'
        for key, value in matched:
            html += f'<div class="item matched"><div class="key">{_esc(key)}</div><div class="value-block">{_esc(value)}</div></div>'
        html += "</div>"

    if differences:
        html += '<div class="section"><h2>有差异</h2>'
        for key, word_value, latex_value in differences:
            html += (
                f'<div class="item difference"><div class="key">{_esc(key)}</div>'
                f'<div class="value-block"><span class="label">Word</span>{_esc(word_value)}</div>'
                f'<div class="value-block"><span class="label">LaTeX</span>{_esc(latex_value)}</div>'
                f"</div>"
            )
        html += "</div>"

    if only_in_one:
        html += '<div class="section"><h2>仅在一方</h2>'
        for source, key, value in only_in_one:
            source_label = "Word" if source == "word" else "LaTeX"
            html += (
                f'<div class="item only"><div class="key">{_esc(key)}</div>'
                f'<div class="value-block"><span class="label">来源</span>{_esc(source_label)}</div>'
                f'<div class="value-block">{_esc(value)}</div>'
                f"</div>"
            )
        html += "</div>"

    html += "</body></html>"
    return html


def render_formatted_text_html(fragments: List[Dict]) -> str:
    """
    将格式片段渲染为 HTML

    Args:
        fragments: 格式片段列表

    Returns:
        HTML 字符串，加粗文本用 <b> 标签
    """
    html_parts = []
    for frag in fragments:
        text = frag["text"]
        # HTML 转义
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if frag.get("bold"):
            html_parts.append(f'<b>{text}</b>')
        else:
            html_parts.append(text)
    return ''.join(html_parts)


def generate_html_report_with_format(matched: List, text_diff: List, format_diff: List, only_in_one: List,
                                     word_file: Path, latex_file: Path) -> str:
    """生成 HTML 格式报告（包含格式对比）"""
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>标题文字对比报告</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 30px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }}
        .stat-card {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
        }}
        .stat-card h3 {{
            margin: 0 0 10px 0;
            font-size: 14px;
            color: #666;
        }}
        .stat-card .value {{
            font-size: 32px;
            font-weight: bold;
        }}
        .matched .value {{ color: #10b981; }}
        .differences .value {{ color: #f59e0b; }}
        .only .value {{ color: #ef4444; }}
        .section {{
            background: white;
            padding: 25px;
            border-radius: 8px;
            margin-bottom: 20px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .section h2 {{
            margin-top: 0;
            padding-bottom: 15px;
            border-bottom: 2px solid #e5e7eb;
        }}
        .item {{
            padding: 15px;
            margin-bottom: 15px;
            border-left: 4px solid #ddd;
            background: #f9fafb;
            border-radius: 4px;
        }}
        .item.matched {{
            border-left-color: #10b981;
            background: #f0fdf4;
        }}
        .item.difference {{
            border-left-color: #f59e0b;
            background: #fffbeb;
        }}
        .item.only {{
            border-left-color: #ef4444;
            background: #fef2f2;
        }}
        .key {{
            font-weight: bold;
            color: #1f2937;
            margin-bottom: 5px;
        }}
        .value {{
            color: #4b5563;
        }}
        .diff-pair {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            margin-top: 10px;
        }}
        .diff-box {{
            padding: 10px;
            background: white;
            border-radius: 4px;
            border: 1px solid #e5e7eb;
        }}
        .diff-box.word {{
            border-left: 3px solid #3b82f6;
        }}
        .diff-box.latex {{
            border-left: 3px solid #8b5cf6;
        }}
        .label {{
            font-size: 12px;
            color: #6b7280;
            margin-bottom: 5px;
        }}
        .meta {{
            color: #9ca3af;
            font-size: 14px;
            margin-top: 30px;
            text-align: center;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>📋 标题文字对比报告</h1>
        <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>

    <div class="stats">
        <div class="stat-card matched">
            <h3>✅ 完全匹配</h3>
            <div class="value">{len(matched)}</div>
        </div>
        <div class="stat-card differences">
            <h3>⚠️ 有差异</h3>
            <div class="value">{len(differences)}</div>
        </div>
        <div class="stat-card only">
            <h3>❌ 仅在一方</h3>
            <div class="value">{len(only_in_one)}</div>
        </div>
    </div>
'''

    # 完全匹配的标题
    if matched:
        html += '<div class="section"><h2>✅ 完全匹配的标题</h2>'
        for key, value in matched:
            html += f'''
    <div class="item matched">
        <div class="key">{key}</div>
        <div class="value">{value}</div>
    </div>'''
        html += '</div>'

    # 有差异的标题
    if differences:
        html += '<div class="section"><h2>⚠️ 有差异的标题</h2>'
        for key, word_value, latex_value in differences:
            html += f'''
    <div class="item difference">
        <div class="key">{key}</div>
        <div class="diff-pair">
            <div class="diff-box word">
                <div class="label">Word 模板</div>
                <div class="value">{word_value}</div>
            </div>
            <div class="diff-box latex">
                <div class="label">LaTeX 文件</div>
                <div class="value">{latex_value}</div>
            </div>
        </div>
    </div>'''
        html += '</div>'

    # 仅在一方的标题
    if only_in_one:
        html += '<div class="section"><h2>❌ 仅在一方的标题</h2>'
        for source, key, value in only_in_one:
            source_label = 'Word 模板' if source == 'word' else 'LaTeX 文件'
            html += f'''
    <div class="item only">
        <div class="key">仅在 {source_label}: {key}</div>
        <div class="value">{value}</div>
    </div>'''
        html += '</div>'

    html += f'''
    <div class="meta">
        <p>Word 文件: {word_file.name}</p>
        <p>LaTeX 文件: {latex_file.name}</p>
    </div>
</body>
</html>'''

    return html


def generate_html_report_with_format(matched: List, text_diff: List, format_diff: List, only_in_one: List,
                                     word_file: Path, latex_file: Path) -> str:
    """生成 HTML 格式报告（包含格式对比）"""
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>标题文字对比报告（含格式）</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            max-width: 1400px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 30px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }}
        .stat-card {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
        }}
        .stat-card h3 {{
            margin: 0 0 10px 0;
            font-size: 14px;
            color: #666;
        }}
        .stat-card .value {{
            font-size: 32px;
            font-weight: bold;
        }}
        .matched .value {{ color: #10b981; }}
        .text-diff .value {{ color: #f59e0b; }}
        .format-diff .value {{ color: #f97316; }}
        .only .value {{ color: #ef4444; }}
        .section {{
            background: white;
            padding: 25px;
            border-radius: 8px;
            margin-bottom: 20px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .section h2 {{
            margin-top: 0;
            padding-bottom: 15px;
            border-bottom: 2px solid #e5e7eb;
        }}
        .item {{
            padding: 15px;
            margin-bottom: 15px;
            border-left: 4px solid #ddd;
            background: #f9fafb;
            border-radius: 4px;
        }}
        .item.matched {{
            border-left-color: #10b981;
            background: #f0fdf4;
        }}
        .item.text-diff {{
            border-left-color: #f59e0b;
            background: #fffbeb;
        }}
        .item.format-diff {{
            border-left-color: #f97316;
            background: #fff7ed;
        }}
        .item.only {{
            border-left-color: #ef4444;
            background: #fef2f2;
        }}
        .key {{
            font-weight: bold;
            color: #1f2937;
            margin-bottom: 8px;
            font-size: 14px;
        }}
        .diff-pair {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            margin-top: 10px;
        }}
        .diff-box {{
            padding: 12px;
            background: white;
            border-radius: 4px;
            border: 1px solid #e5e7eb;
        }}
        .diff-box.word {{
            border-left: 3px solid #3b82f6;
        }}
        .diff-box.latex {{
            border-left: 3px solid #8b5cf6;
        }}
        .label {{
            font-size: 12px;
            color: #6b7280;
            margin-bottom: 8px;
            font-weight: 500;
        }}
        .rendered-text {{
            font-size: 15px;
            line-height: 1.8;
            color: #1f2937;
        }}
        .rendered-text b {{
            font-weight: 700;
            color: #1e3a8a;
        }}
        .diff-marker {{
            margin-top: 12px;
            padding: 10px;
            background: #fef3c7;
            border-radius: 4px;
            font-size: 13px;
            color: #92400e;
        }}
        .diff-marker-item {{
            padding: 4px 0;
            border-bottom: 1px solid #fde68a;
        }}
        .diff-marker-item:last-child {{
            border-bottom: none;
        }}
        .meta {{
            color: #9ca3af;
            font-size: 14px;
            margin-top: 30px;
            text-align: center;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>📋 标题文字对比报告（含格式）</h1>
        <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>

    <div class="stats">
        <div class="stat-card matched">
            <h3>✅ 完全匹配</h3>
            <div class="value">{len(matched)}</div>
        </div>
        <div class="stat-card text-diff">
            <h3>⚠️ 文本差异</h3>
            <div class="value">{len(text_diff)}</div>
        </div>
        <div class="stat-card format-diff">
            <h3>🔶 格式差异</h3>
            <div class="value">{len(format_diff)}</div>
        </div>
        <div class="stat-card only">
            <h3>❌ 仅在一方</h3>
            <div class="value">{len(only_in_one)}</div>
        </div>
    </div>
'''

    # 完全匹配的标题
    if matched:
        html += '<div class="section"><h2>✅ 完全匹配的标题（文本+格式）</h2>'
        for key, text, result in matched:
            word_html = render_formatted_text_html(result.get("word_fragments", []))
            html += f'''
    <div class="item matched">
        <div class="key">{key}</div>
        <div class="rendered-text">{word_html}</div>
    </div>'''
        html += '</div>'

    # 文本差异
    if text_diff:
        html += '<div class="section"><h2>⚠️ 文本差异</h2>'
        for key, word_value, latex_value in text_diff:
            html += f'''
    <div class="item text-diff">
        <div class="key">{key}</div>
        <div class="diff-pair">
            <div class="diff-box word">
                <div class="label">Word 模板</div>
                <div class="rendered-text">{word_value}</div>
            </div>
            <div class="diff-box latex">
                <div class="label">LaTeX 文件</div>
                <div class="rendered-text">{latex_value}</div>
            </div>
        </div>
    </div>'''
        html += '</div>'

    # 格式差异
    if format_diff:
        html += '<div class="section"><h2>🔶 格式差异（加粗）</h2>'
        for key, text, result in format_diff:
            word_html = render_formatted_text_html(result.get("word_fragments", []))
            latex_html = render_formatted_text_html(result.get("latex_fragments", []))

            # 构建差异标记
            diff_markers = []
            for diff in result.get("differences", []):
                char = diff.get("char", "")
                pos = diff.get("position", 0)
                word_bold = "加粗" if diff.get("word_bold") else "正常"
                latex_bold = "加粗" if diff.get("latex_bold") else "正常"
                diff_markers.append(f'位置 {pos}: "{char}" - Word:{word_bold}, LaTeX:{latex_bold}')

            diff_marker_html = ""
            if diff_markers:
                diff_marker_html = '<div class="diff-marker">' + \
                    ''.join(f'<div class="diff-marker-item">{marker}</div>' for marker in diff_markers) + \
                    '</div>'

            html += f'''
    <div class="item format-diff">
        <div class="key">{key}</div>
        <div class="diff-pair">
            <div class="diff-box word">
                <div class="label">Word 模板</div>
                <div class="rendered-text">{word_html}</div>
            </div>
            <div class="diff-box latex">
                <div class="label">LaTeX 文件</div>
                <div class="rendered-text">{latex_html}</div>
            </div>
        </div>
        {diff_marker_html}
    </div>'''
        html += '</div>'

    # 仅在一方的标题
    if only_in_one:
        html += '<div class="section"><h2>❌ 仅在一方的标题</h2>'
        for source, key, value in only_in_one:
            source_label = 'Word 模板' if source == 'word' else 'LaTeX 文件'
            html += f'''
    <div class="item only">
        <div class="key">仅在 {source_label}: {key}</div>
        <div class="rendered-text">{value}</div>
    </div>'''
        html += '</div>'

    html += f'''
    <div class="meta">
        <p>Word 文件: {word_file.name}</p>
        <p>LaTeX 文件: {latex_file.name}</p>
    </div>
</body>
</html>'''

    return html


def generate_latex_fix_suggestions(format_diff: List) -> str:
    """
    生成 LaTeX 修复建议

    Args:
        format_diff: 格式差异列表

    Returns:
        LaTeX 修复代码字符串
    """
    lines = []
    lines.append('% LaTeX 标题格式修复建议')
    lines.append('% 自动生成于: ' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    lines.append('% 请根据实际情况修改 main.tex 中的对应标题')
    lines.append('')
    lines.append('% 使用方法：')
    lines.append('% 1. 将下面的 \\section{} 或 \\subsection{} 替换到 main.tex 中')
    lines.append('% 2. 确保格式符合 Word 模板要求')
    lines.append('')

    if not format_diff:
        lines.append('% ✅ 所有标题格式一致，无需修复')
        return '\n'.join(lines)

    lines.append('% 修复建议：')
    lines.append('')

    for key, text, result in format_diff:
        word_fragments = result.get("word_fragments", [])

        # 生成 LaTeX 代码
        latex_parts = []
        for frag in word_fragments:
            frag_text = frag["text"]
            if frag.get("bold"):
                latex_parts.append(f'\\textbf{{{frag_text}}}')
            else:
                latex_parts.append(frag_text)

        latex_code = ''.join(latex_parts)

        # 判断是 section 还是 subsection
        if key.startswith('section_'):
            command = '\\section'
        elif key.startswith('subsection_'):
            command = '\\subsection'
        else:
            command = '\\section'

        lines.append(f'% {key}: {text}')
        lines.append(f'{command}{{{latex_code}}}')
        lines.append('')

    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='对比基准模板（推荐 PDF）与 LaTeX 的标题文字')
    parser.add_argument('source_file', type=Path, help='基准文件路径（推荐: .pdf；兼容: .docx）')
    parser.add_argument('latex_file', type=Path, help='LaTeX 文件路径（main.tex）')
    parser.add_argument('--report', type=Path, help='输出报告文件路径')
    parser.add_argument('--format', choices=['auto', 'text', 'html'], default='auto',
                       help='报告格式（auto 根据扩展名自动判断）')
    parser.add_argument('--check-format', action='store_true',
                       help='检查格式（加粗）是否一致（默认仅检查文本）')
    parser.add_argument('--fix-file', type=Path, help='输出 LaTeX 修复建议文件路径')

    args = parser.parse_args()

    # 提取标题
    print(f'📖 正在提取基准标题: {args.source_file}')
    word_headings = extract_from_source(args.source_file, check_format=args.check_format)

    print(f'📖 正在提取 LaTeX 标题: {args.latex_file}')
    latex_headings = extract_from_latex(args.latex_file, check_format=args.check_format)

    # 对比标题
    if args.check_format:
        print('🔍 正在对比标题（包含格式）...')
        matched, text_diff, format_diff, only_in_one = compare_headings_with_format(
            word_headings, latex_headings
        )
    else:
        print('🔍 正在对比标题...')
        matched, differences, only_in_one = compare_headings(word_headings, latex_headings)
        text_diff = []
        format_diff = []
        # 将旧的 differences 转换为 text_diff 格式以保持一致性
        text_diff = differences

    # 生成报告
    if args.report:
        # 判断格式
        if args.format == 'auto':
            if args.report.suffix == '.html':
                fmt = 'html'
            elif args.report.suffix == '.md':
                fmt = 'markdown'
            else:
                fmt = 'text'
        else:
            fmt = args.format

        if args.check_format:
            # 格式对比模式
            if fmt == 'html':
                report = generate_html_report_with_format(matched, text_diff, format_diff, only_in_one,
                                                          args.source_file, args.latex_file)
            else:
                report = generate_text_report_with_format(matched, text_diff, format_diff, only_in_one)
        else:
            # 传统模式
            if fmt == 'html':
                report = generate_html_report(matched, differences, only_in_one,
                                            args.source_file, args.latex_file)
            else:
                report = generate_text_report(matched, differences, only_in_one)

        with open(args.report, 'w', encoding='utf-8') as f:
            f.write(report)

        if args.check_format:
            total = len(matched) + len(text_diff) + len(format_diff)
            print(f'✅ 报告已生成: {args.report}')
            print(f'   总计: {total} | 匹配: {len(matched)} | 文本差异: {len(text_diff)} | 格式差异: {len(format_diff)} | 仅在一方: {len(only_in_one)}')
        else:
            print(f'✅ 报告已生成: {args.report}')
            print(f'   总计: {len(matched) + len(differences)} | 匹配: {len(matched)} | 差异: {len(differences)} | 仅在一方: {len(only_in_one)}')

    else:
        # 打印到控制台
        if args.check_format:
            report = generate_text_report_with_format(matched, text_diff, format_diff, only_in_one)
        else:
            report = generate_text_report(matched, differences, only_in_one)
        print(report)

    # 生成修复建议文件
    if args.fix_file and args.check_format and format_diff:
        fix_content = generate_latex_fix_suggestions(format_diff)
        with open(args.fix_file, 'w', encoding='utf-8') as f:
            f.write(fix_content)
        print(f'🔧 LaTeX 修复建议已生成: {args.fix_file}')
    elif args.fix_file and args.check_format and not format_diff:
        # 无格式差异，仍然生成文件
        fix_content = generate_latex_fix_suggestions(format_diff)
        with open(args.fix_file, 'w', encoding='utf-8') as f:
            f.write(fix_content)
        print(f'✅ 所有标题格式一致，修复建议文件已生成: {args.fix_file}')


if __name__ == '__main__':
    main()
